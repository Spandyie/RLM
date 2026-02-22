"""
Document Processor - Parse and chunk documents.

Supports: PDF, TXT, MD, DOCX
"""

import os
import io
import hashlib
from dataclasses import dataclass
from typing import List, Optional

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


@dataclass
class DocumentChunk:
    """A piece of a document."""
    text: str
    doc_id: str
    chunk_index: int
    metadata: dict


@dataclass
class ProcessedDocument:
    """A processed document with chunks."""
    doc_id: str
    filename: str
    chunks: List[DocumentChunk]
    full_text: str
    metadata: dict


class DocumentProcessor:
    """Parse and chunk documents."""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    async def process(
        self,
        file_content: bytes,
        filename: str,
        doc_id: Optional[str] = None,
    ) -> ProcessedDocument:
        """Process a document into chunks."""

        if doc_id is None:
            doc_id = hashlib.md5(file_content[:1000]).hexdigest()[:12]

        ext = os.path.splitext(filename)[1].lower()

        if ext == ".pdf":
            text = self._extract_pdf(file_content)
        elif ext == ".docx":
            text = self._extract_docx(file_content)
        else:
            text = file_content.decode("utf-8", errors="ignore")

        text = self._clean(text)
        chunks = self._chunk(text, doc_id, filename)

        return ProcessedDocument(
            doc_id=doc_id,
            filename=filename,
            chunks=chunks,
            full_text=text,
            metadata={"filename": filename, "chunks": len(chunks)},
        )

    def _extract_pdf(self, content: bytes) -> str:
        if PdfReader is None:
            raise ImportError("Install pypdf2: pip install pypdf2")

        reader = PdfReader(io.BytesIO(content))
        parts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                parts.append(f"[Page {i+1}]\n{text}")
        return "\n\n".join(parts)

    def _extract_docx(self, content: bytes) -> str:
        if DocxDocument is None:
            raise ImportError("Install python-docx: pip install python-docx")

        doc = DocxDocument(io.BytesIO(content))
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def _clean(self, text: str) -> str:
        import re
        text = re.sub(r" +", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def _chunk(self, text: str, doc_id: str, filename: str) -> List[DocumentChunk]:
        chunks = []
        paragraphs = text.split("\n\n")

        current = ""
        idx = 0

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            if len(current) + len(para) > self.chunk_size:
                if current:
                    chunks.append(DocumentChunk(
                        text=current.strip(),
                        doc_id=doc_id,
                        chunk_index=idx,
                        metadata={"filename": filename},
                    ))
                    idx += 1
                    # Add overlap
                    overlap = current[-self.chunk_overlap:] if len(current) > self.chunk_overlap else ""
                    current = overlap + "\n\n" + para if overlap else para
                else:
                    current = para
            else:
                current = current + "\n\n" + para if current else para

        if current.strip():
            chunks.append(DocumentChunk(
                text=current.strip(),
                doc_id=doc_id,
                chunk_index=idx,
                metadata={"filename": filename},
            ))

        return chunks
